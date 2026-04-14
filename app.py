from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import HTMLResponse, FileResponse
from docx import Document
from tempfile import NamedTemporaryFile
import shutil
import os
import re
import html

app = FastAPI(title="Rellenador de prácticas de laboratorio ACE3 - IPN")


# =========================
# Lógica de relleno DOCX
# =========================

def normalizar_etiqueta(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip().upper())


def limpiar_runs(paragraph):
    if not paragraph.runs:
        return None
    first = paragraph.runs[0]
    for r in paragraph.runs[1:]:
        r.text = ""
    return first


def set_paragraph_text(paragraph, new_text: str):
    first = limpiar_runs(paragraph)
    if first is None:
        paragraph.add_run(new_text)
    else:
        first.text = new_text


def construir_mapa_simple(datos: dict) -> dict:
    personales = datos.get("datos_personales", {})
    extra = datos.get("campos_extra", {})
    overrides = datos.get("mapeo_etiquetas_personalizado", {})

    mapa = {
        "NOMBRE": personales.get("nombre", ""),
        "BOLETA": personales.get("boleta", ""),
        "GRUPO": personales.get("grupo", ""),
        "SUBGRUPO": personales.get("subgrupo", ""),
        "SECCIÓN": personales.get("seccion", ""),
        "SECCION": personales.get("seccion", ""),
        "FECHA": personales.get("fecha", ""),
        "PROFESOR": extra.get("profesor", ""),
        "PROFESORES": extra.get("profesor", ""),
        "FIRMA PROF": extra.get("firma_prof", ""),
        "FIRMA PROF.": extra.get("firma_prof", ""),
        "ING.": extra.get("ingeniero", ""),
        "CALIFICACIÓN": extra.get("calificacion", ""),
        "CALIFICACION": extra.get("calificacion", ""),
    }

    for k, v in overrides.items():
        mapa[k] = v

    return mapa


def reemplazar_si_es_etiqueta_exacta(texto: str, mapa_simple: dict):
    texto_limpio = texto.strip()
    texto_norm = normalizar_etiqueta(texto_limpio).rstrip(":")

    if texto_norm in mapa_simple and mapa_simple[texto_norm] != "":
        valor = str(mapa_simple[texto_norm])
        if texto_limpio.endswith(":"):
            return f"{texto_limpio} {valor}"
        return f"{texto_limpio}\n{valor}"

    return None


def reemplazar_etiquetas_inline(texto: str, mapa_simple: dict):
    nuevo = texto
    for etiqueta, valor in mapa_simple.items():
        if valor == "":
            continue
        patron = rf"(?<!\w)({re.escape(etiqueta)}\s*:)"
        nuevo = re.sub(patron, rf"\1 {valor}", nuevo, flags=re.IGNORECASE)
    return nuevo


def recorrer_tablas(tablas):
    for table in tablas:
        yield table
        for row in table.rows:
            for cell in row.cells:
                for inner in recorrer_tablas(cell.tables):
                    yield inner


def procesar_parrafos_de_celda(cell, mapa_simple):
    for p in cell.paragraphs:
        original = p.text
        if not original.strip():
            continue

        exacto = reemplazar_si_es_etiqueta_exacta(original, mapa_simple)
        if exacto is not None:
            set_paragraph_text(p, exacto)
            continue

        inline = reemplazar_etiquetas_inline(original, mapa_simple)
        if inline != original:
            set_paragraph_text(p, inline)


def procesar_documento(doc, mapa_simple):
    for p in doc.paragraphs:
        original = p.text
        if not original.strip():
            continue

        exacto = reemplazar_si_es_etiqueta_exacta(original, mapa_simple)
        if exacto is not None:
            set_paragraph_text(p, exacto)
            continue

        inline = reemplazar_etiquetas_inline(original, mapa_simple)
        if inline != original:
            set_paragraph_text(p, inline)

    for table in recorrer_tablas(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                texto_celda = cell.text.strip()

                exacto = reemplazar_si_es_etiqueta_exacta(texto_celda, mapa_simple)
                if exacto is not None:
                    if cell.paragraphs:
                        set_paragraph_text(cell.paragraphs[0], exacto)
                        for extra_p in cell.paragraphs[1:]:
                            set_paragraph_text(extra_p, "")
                    else:
                        cell.text = exacto
                else:
                    procesar_parrafos_de_celda(cell, mapa_simple)


def aplicar_relleno_por_coordenadas(doc, datos):
    for item in datos.get("relleno_celdas", []):
        tabla_idx = item["tabla"]
        fila = item["fila"]
        columna = item["columna"]
        valor = str(item.get("valor", ""))
        modo = item.get("modo", "reemplazar")

        if tabla_idx >= len(doc.tables):
            continue

        tabla = doc.tables[tabla_idx]

        if fila >= len(tabla.rows) or columna >= len(tabla.rows[fila].cells):
            continue

        cell = tabla.rows[fila].cells[columna]
        actual = cell.text.strip()

        if modo == "anexar" and actual:
            nuevo = f"{actual}\n{valor}"
        else:
            nuevo = valor

        if cell.paragraphs:
            set_paragraph_text(cell.paragraphs[0], nuevo)
            for extra_p in cell.paragraphs[1:]:
                set_paragraph_text(extra_p, "")
        else:
            cell.text = nuevo


def rellenar_docx(docx_entrada: str, datos: dict, docx_salida: str):
    doc = Document(docx_entrada)

    mapa_simple = construir_mapa_simple(datos)
    procesar_documento(doc, mapa_simple)
    aplicar_relleno_por_coordenadas(doc, datos)

    doc.save(docx_salida)


# =========================
# Interfaz HTML sencilla
# =========================

def pagina_base(contenido: str) -> str:
    return f"""
    <!doctype html>
    <html lang="es">
    <head>
        <meta charset="utf-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Rellenador de prácticas de laboratorio ACE3 - IPN</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                background: #f4f6fb;
                margin: 0;
                padding: 24px;
            }}
            .contenedor {{
                max-width: 780px;
                margin: 0 auto;
                background: white;
                padding: 28px;
                border-radius: 16px;
                box-shadow: 0 8px 30px rgba(0,0,0,0.08);
            }}
            h1 {{
                margin-top: 0;
                color: #17324d;
            }}
            h2 {{
                color: #244b73;
            }}
            .grid {{
                display: grid;
                grid-template-columns: 1fr 1fr;
                gap: 16px;
            }}
            .full {{
                grid-column: 1 / -1;
            }}
            label {{
                display: block;
                margin-bottom: 6px;
                font-weight: bold;
            }}
            input {{
                width: 100%;
                padding: 12px;
                border: 1px solid #cfd8e3;
                border-radius: 10px;
                box-sizing: border-box;
            }}
            button {{
                background: #0d6efd;
                color: white;
                border: 0;
                padding: 12px 18px;
                border-radius: 10px;
                cursor: pointer;
                font-size: 16px;
            }}
            .secundario {{
                background: #6c757d;
            }}
            .resumen {{
                background: #f8fafc;
                border: 1px solid #dbe4ee;
                border-radius: 12px;
                padding: 16px;
                margin: 16px 0;
            }}
            .nota {{
                color: #5d6b79;
                font-size: 14px;
            }}
            @media (max-width: 640px) {{
                .grid {{
                    grid-template-columns: 1fr;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="contenedor">
            {contenido}
        </div>
    </body>
    </html>
    """


@app.get("/", response_class=HTMLResponse)
async def inicio():
    contenido = """
    <h1>Rellenador de prácticas de laboratorio ACE3 - IPN</h1>
    <p class="nota">Llena tus datos, sube tu archivo .docx y descarga tu práctica rellenada.</p>

    <form action="/confirmar" method="post" enctype="multipart/form-data">
        <div class="grid">
            <div>
                <label>Nombre completo</label>
                <input name="nombre" required>
            </div>

            <div>
                <label>Boleta</label>
                <input name="boleta" required>
            </div>

            <div>
                <label>Profesor</label>
                <input name="profesor" required>
            </div>

            <div>
                <label>Fecha</label>
                <input type="date" name="fecha" required>
            </div>

            <div>
                <label>Grupo</label>
                <input name="grupo" value="6EV1">
            </div>

            <div>
                <label>Subgrupo</label>
                <input name="subgrupo" value="1">
            </div>

            <div>
                <label>Sección</label>
                <input name="seccion" value="A">
            </div>

            <div class="full">
                <label>Archivo Word (.docx)</label>
                <input type="file" name="archivo" accept=".docx" required>
            </div>
        </div>

        <br>
        <button type="submit">Continuar</button>
    </form>
    """
    return pagina_base(contenido)


@app.post("/confirmar", response_class=HTMLResponse)
async def confirmar(
    nombre: str = Form(...),
    boleta: str = Form(...),
    profesor: str = Form(...),
    fecha: str = Form(...),
    grupo: str = Form(""),
    subgrupo: str = Form(""),
    seccion: str = Form(""),
    archivo: UploadFile = File(...)
):
    if not archivo.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Solo se permiten archivos .docx")

    temporal_entrada = NamedTemporaryFile(delete=False, suffix=".docx")
    with temporal_entrada as f:
        shutil.copyfileobj(archivo.file, f)

    esc = html.escape
    esc_attr = lambda x: html.escape(x, quote=True)

    contenido = f"""
    <h1>Rellenador de prácticas de laboratorio ACE3 - IPN</h1>
    <h2>Confirma tus datos</h2>

    <div class="resumen">
        <p><strong>Nombre:</strong> {esc(nombre)}</p>
        <p><strong>Boleta:</strong> {esc(boleta)}</p>
        <p><strong>Profesor:</strong> {esc(profesor)}</p>
        <p><strong>Fecha:</strong> {esc(fecha)}</p>
        <p><strong>Grupo:</strong> {esc(grupo)}</p>
        <p><strong>Subgrupo:</strong> {esc(subgrupo)}</p>
        <p><strong>Sección:</strong> {esc(seccion)}</p>
    </div>

    <p>¿Los datos son correctos?</p>

    <form action="/procesar" method="post">
        <input type="hidden" name="nombre" value="{esc_attr(nombre)}">
        <input type="hidden" name="boleta" value="{esc_attr(boleta)}">
        <input type="hidden" name="profesor" value="{esc_attr(profesor)}">
        <input type="hidden" name="fecha" value="{esc_attr(fecha)}">
        <input type="hidden" name="grupo" value="{esc_attr(grupo)}">
        <input type="hidden" name="subgrupo" value="{esc_attr(subgrupo)}">
        <input type="hidden" name="seccion" value="{esc_attr(seccion)}">
        <input type="hidden" name="tmp_path" value="{esc_attr(temporal_entrada.name)}">
        <button type="submit">Sí, generar archivo</button>
    </form>

    <br>

    <form action="/editar" method="post">
        <input type="hidden" name="nombre" value="{esc_attr(nombre)}">
        <input type="hidden" name="boleta" value="{esc_attr(boleta)}">
        <input type="hidden" name="profesor" value="{esc_attr(profesor)}">
        <input type="hidden" name="fecha" value="{esc_attr(fecha)}">
        <input type="hidden" name="grupo" value="{esc_attr(grupo)}">
        <input type="hidden" name="subgrupo" value="{esc_attr(subgrupo)}">
        <input type="hidden" name="seccion" value="{esc_attr(seccion)}">
        <input type="hidden" name="tmp_path" value="{esc_attr(temporal_entrada.name)}">
        <button type="submit" class="secundario">No, editar datos</button>
    </form>
    """

    return pagina_base(contenido)


@app.post("/editar", response_class=HTMLResponse)
async def editar(
    nombre: str = Form(...),
    boleta: str = Form(...),
    profesor: str = Form(...),
    fecha: str = Form(...),
    grupo: str = Form(""),
    subgrupo: str = Form(""),
    seccion: str = Form(""),
    tmp_path: str = Form(...)
):
    esc_attr = lambda x: html.escape(x, quote=True)

    contenido = f"""
    <h1>Rellenador de prácticas de laboratorio ACE3 - IPN</h1>
    <h2>Corrige tus datos</h2>

    <form action="/procesar" method="post">
        <div class="grid">
            <div>
                <label>Nombre completo</label>
                <input name="nombre" value="{esc_attr(nombre)}" required>
            </div>

            <div>
                <label>Boleta</label>
                <input name="boleta" value="{esc_attr(boleta)}" required>
            </div>

            <div>
                <label>Profesor</label>
                <input name="profesor" value="{esc_attr(profesor)}" required>
            </div>

            <div>
                <label>Fecha</label>
                <input name="fecha" value="{esc_attr(fecha)}" required>
            </div>

            <div>
                <label>Grupo</label>
                <input name="grupo" value="{esc_attr(grupo)}">
            </div>

            <div>
                <label>Subgrupo</label>
                <input name="subgrupo" value="{esc_attr(subgrupo)}">
            </div>

            <div>
                <label>Sección</label>
                <input name="seccion" value="{esc_attr(seccion)}">
            </div>
        </div>

        <input type="hidden" name="tmp_path" value="{esc_attr(tmp_path)}">

        <br>
        <button type="submit">Generar archivo</button>
    </form>
    """

    return pagina_base(contenido)


@app.post("/procesar")
async def procesar(
    nombre: str = Form(...),
    boleta: str = Form(...),
    profesor: str = Form(...),
    fecha: str = Form(...),
    grupo: str = Form(""),
    subgrupo: str = Form(""),
    seccion: str = Form(""),
    tmp_path: str = Form(...)
):
    if not os.path.exists(tmp_path):
        raise HTTPException(status_code=400, detail="El archivo temporal ya no existe. Vuelve a cargar el documento.")

    datos = {
        "datos_personales": {
            "nombre": nombre,
            "boleta": boleta,
            "grupo": grupo,
            "subgrupo": subgrupo,
            "seccion": seccion,
            "fecha": fecha
        },
        "campos_extra": {
            "profesor": profesor,
            "firma_prof": "",
            "ingeniero": "",
            "calificacion": ""
        },
        "mapeo_etiquetas_personalizado": {
            "NOMBRE": nombre,
            "BOLETA": boleta,
            "PROFESOR": profesor,
            "PROFESORES": profesor,
            "FECHA": fecha,
            "GRUPO": grupo,
            "SUBGRUPO": subgrupo,
            "SECCION": seccion,
            "SECCIÓN": seccion
        },
        "relleno_celdas": []
    }

    salida = NamedTemporaryFile(delete=False, suffix="_rellenado.docx")
    salida.close()

    rellenar_docx(tmp_path, datos, salida.name)

    nombre_descarga = f"practica_{boleta}_rellenada.docx"

    return FileResponse(
        path=salida.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=nombre_descarga
    )
